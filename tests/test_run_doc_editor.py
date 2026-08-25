"""The in-app editor safely edits every recognised run-doc section."""
from pathlib import Path

import pytest
from docx import Document

import paths
import run_doc_editor as editor


def _make_doc(path: Path):
    doc = Document()
    doc.add_paragraph("SERVPRO DAILY RUN")
    doc.add_paragraph("Date: 8/24/2026")
    doc.add_paragraph("Monitor")
    first = doc.add_paragraph("Alpha: monitor")
    first.runs[0].font.strike = True
    doc.add_paragraph("Bravo: equipment pickup")
    doc.add_paragraph("")
    doc.add_paragraph("Work To Be Performed")
    doc.add_paragraph("Charlie: demo")
    doc.add_paragraph("Warehouse (Charlie): Maria")
    doc.add_paragraph("Delta: new loss")
    doc.add_paragraph("")
    doc.add_paragraph("Upcoming")
    doc.add_paragraph("Do not touch this future row")
    doc.save(path)


@pytest.fixture
def run_doc_file(tmp_path, monkeypatch):
    monkeypatch.setattr(paths, "data", lambda name: str(tmp_path / "app" / name))
    path = tmp_path / "Monday 8.24.26.docx"
    _make_doc(path)
    return path


def test_load_includes_struck_and_warehouse_rows(run_doc_file):
    model = editor.read_document(str(run_doc_file))
    assert [row["text"] for row in model["sections"]["monitor"]] == [
        "Alpha: monitor", "Bravo: equipment pickup"]
    assert model["sections"]["monitor"][0]["struck"] is True
    assert [row["text"] for row in model["sections"]["work"]] == [
        "Charlie: demo", "Warehouse (Charlie): Maria", "Delta: new loss"]
    assert [row["text"] for row in model["sections"]["upcoming"]] == [
        "Do not touch this future row"]
    assert model["section_order"] == ["monitor", "work", "upcoming"]
    assert editor.SECTION_LABELS["pending_property"].startswith(
        "Pending Approvals")


def test_save_preserves_authored_spacing(run_doc_file):
    model = editor.read_document(str(run_doc_file))
    model["sections"]["monitor"][0]["text"] = "Alpha:  monitor   9-11 "
    saved = editor.save_document(
        str(run_doc_file), model["version"], model["sections"])
    assert saved["sections"]["monitor"][0]["text"] == "Alpha:  monitor   9-11 "


def test_save_reorders_adds_deletes_and_preserves_the_rest(run_doc_file):
    before = editor.read_document(str(run_doc_file))
    sections = {
        "monitor": [
            {"text": "Bravo: equipment pickup", "struck": False},
            {"text": "Alpha: monitor", "struck": False},
            {"text": "Echo: final monitor", "struck": True},
        ],
        "work": [
            {"text": "Delta: new loss", "struck": False},
            {"text": "Charlie: revised demo", "struck": True},
        ],
    }
    saved = editor.save_document(str(run_doc_file), before["version"], sections)
    got = {name: [{"text": row["text"], "struck": row["struck"]}
                  for row in saved["sections"][name]]
           for name in sections}
    assert got == sections
    text = [paragraph.text for paragraph in Document(run_doc_file).paragraphs]
    assert "Upcoming" in text
    assert "Do not touch this future row" in text
    assert "Warehouse (Charlie): Maria" not in text
    assert Path(saved["backup"]).is_file()


def test_stale_screen_cannot_overwrite_a_newer_word_copy(run_doc_file):
    model = editor.read_document(str(run_doc_file))
    doc = Document(run_doc_file)
    doc.paragraphs[-1].text = "Changed outside Linguar Hub"
    doc.save(run_doc_file)
    with pytest.raises(editor.RunDocConflict):
        editor.save_document(str(run_doc_file), model["version"], model["sections"])


def test_failed_validation_never_replaces_the_original(run_doc_file, monkeypatch):
    model = editor.read_document(str(run_doc_file))
    original = run_doc_file.read_bytes()
    monkeypatch.setattr(editor, "read_document",
                        lambda path: {"sections": {"monitor": [], "work": []}})
    with pytest.raises(ValueError, match="did not validate"):
        editor.save_document(str(run_doc_file), model["version"], model["sections"])
    assert run_doc_file.read_bytes() == original


def test_editor_is_wired_into_the_work_sidebar():
    import home_web

    work = dict((key, label) for group, items in home_web.NAV_GROUPS
                if group == "Work" for key, _icon, label in items)
    assert work["run_doc_editor"] == "Run Doc Editor"
    assert home_web.SUB_MODULES["run_doc_editor"] == "run_doc_editor_web"


def test_editor_ui_has_drag_keyboard_and_conflict_guardrails():
    root = Path(__file__).resolve().parents[1] / "run_doc_editor_web_assets"
    js = (root / "app.js").read_text(encoding="utf-8")
    html = (root / "index.html").read_text(encoding="utf-8")
    assert "dragstart" in js and "dropRow" in js
    assert "drag-before" in js and "drag-after" in js
    assert "target.index + (before ? 0 : 1)" in js
    assert "Ctrl+S" in html and "event.key.toLowerCase()" in js
    assert "state.model.version" in js and "result?.conflict" in js
    assert 'id="section-rail"' in html
    assert "section_order" in js and "section_labels" in js


def test_editor_has_optional_section_aware_item_template():
    root = Path(__file__).resolve().parents[1] / "run_doc_editor_web_assets"
    js = (root / "app.js").read_text(encoding="utf-8")
    html = (root / "index.html").read_text(encoding="utf-8")
    for field in ("job", "address", "phone", "task", "date", "time",
                  "crew", "status"):
        assert f'id="field-{field}"' in html
    assert "function formatRunItem" in js
    assert "DATED_SECTIONS.has(section)" in js
    assert "This item belongs to the selected day, so no date is needed" in js
    assert "Every field is optional" in js
