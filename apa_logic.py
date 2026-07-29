"""APA Monitor logic — section model, doc parse/write, text normalizers.

UI-free (no tkinter): the pure logic + data the APA Tk panel and the web
panel share. `apa_monitor_gui` re-exports every name below as a backward-
compat shim, so the Tk UI keeps working unchanged and external callers
(apa_web, audit_web, iuq, home_web) are unaffected.

`parse_existing_doc` / `write_doc` take an optional `section_order`; when
omitted they fall back to `_persisted_section_order()` — which is exactly
the value the old module-global `SECTION_ORDER` held — so every existing
caller keeps identical behavior while new callers can inject a custom order.

Extracted from apa_monitor_gui.py — see EMS_Tk_Extraction_Plan.md.
"""
from __future__ import annotations

import os
import re
import urllib.parse
from datetime import datetime, timedelta

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor, Pt, Inches

import config
import paths

# APA Monitor docs root. Resolved lazily from config each access
# (config.load() is mtime-cached) so a Settings change or department (OC/IE)
# switch takes effect without a restart. `apa_logic.APA_ROOT` (and
# `from apa_logic import APA_ROOT`) still resolve via the module __getattr__.
def _apa_root():
    return config.load().get("apa_monitor_root") or r"X:\IE_Public\APA Monitor"


def __getattr__(name):
    if name == "APA_ROOT":
        return _apa_root()
    raise AttributeError(f"module {__name__!r} has no attribute {name!r}")

_ITEM_FONT_SIZE = Pt(9)
_EXTENDED_RED   = RGBColor(0xC0, 0x39, 0x2B)
_NUM_PREFIX_RE  = re.compile(r'^\d+\.\s*')


# ── Section name constants ───────────────────────────────────────────────────
# Edit any name here once and every dropdown / parser / writer follows.
SEC_FINAL_UPLOADS      = "Final Uploads"
SEC_EST_MISSING        = "ESTIMATING MISSING/ADDITIONAL ITEMS"
SEC_EST_SERVICE_CALL   = "ESTIMATING SERVICE CALL"
SEC_EST_TBA            = "ESTIMATING TBA"
SEC_EST_SNAPSHOT       = "ESTIMATING SNAPSHOT"
SEC_PENDING_REVIEW     = "PENDING REVIEW"
SEC_PENDING_REVIEW_DOC = "PENDING REVIEW ."   # how it appears in the .docx
SEC_INITIAL_UPLOADS    = "Initial Uploads"
SEC_DAILY_UPLOADS      = "Daily Uploads"
SEC_AUDIT_REJECTION    = "Audit Rejection"
SEC_AUDIT_DISPUTE      = "Audit Dispute"

# Default estimator roster — the user can add / reorder / remove via the
# "👥 Manage estimators" dialog, which persists the override. Anything saved
# wins over this default; first run with no persistence falls back here.
_DEFAULT_ESTIMATORS_ORDERED = [
    "JUAN", "AARON", "JOHNNY", "KIM", "ZAC",
    "ESTEBAN", "VICTORIA", "PABLO", "AARON L", "SAMANTHA", "RECON",
]

# Built-in sections — always exist; reorderable but not removable.
_BUILTIN_SECTIONS = (
    SEC_FINAL_UPLOADS,
    SEC_EST_MISSING,
    SEC_EST_SERVICE_CALL,
    SEC_EST_TBA,
    SEC_EST_SNAPSHOT,
    SEC_PENDING_REVIEW,
    SEC_INITIAL_UPLOADS,
    SEC_DAILY_UPLOADS,
    SEC_AUDIT_REJECTION,
    SEC_AUDIT_DISPUTE,
)
_BUILTIN_SET = set(_BUILTIN_SECTIONS)

_DEFAULT_SECTION_ORDER = [
    SEC_FINAL_UPLOADS,
    SEC_EST_MISSING,
    SEC_EST_SERVICE_CALL,
    SEC_EST_TBA,
    SEC_EST_SNAPSHOT,
    SEC_PENDING_REVIEW,
    *_DEFAULT_ESTIMATORS_ORDERED,
    SEC_INITIAL_UPLOADS,
    SEC_DAILY_UPLOADS,
    SEC_AUDIT_REJECTION,
    SEC_AUDIT_DISPUTE,
]


def _persisted_section_order():
    """Return the user's persisted full section order, falling back to
    `_DEFAULT_SECTION_ORDER`. Reads the new `apa_section_order` key
    first; if absent, migrates from the legacy `apa_estimators` key
    (which only stored the estimator slice).

    Always ensures every built-in section is present — a saved order
    that's missing one (e.g. a stale persistence file from before a
    code-level addition) gets the missing builtin(s) appended in
    `_DEFAULT_SECTION_ORDER` order, so the panel never silently drops
    a section."""
    try:
        import persistence as _per
        raw = _per.get("apa_section_order")
        if isinstance(raw, list) and raw:
            cleaned = [str(e).strip() for e in raw if str(e).strip()]
            cleaned = [(e.upper() if e not in _BUILTIN_SET else e)
                        for e in cleaned]
            # Make sure no builtin is missing — append any that are.
            missing = [b for b in _BUILTIN_SECTIONS if b not in cleaned]
            return cleaned + missing
        # Legacy: only the estimator slice was persisted. Splice it into
        # the default order in the canonical position.
        legacy = _per.get("apa_estimators")
        if isinstance(legacy, list) and legacy:
            ests = [str(e).strip().upper() for e in legacy if str(e).strip()]
            return [
                SEC_FINAL_UPLOADS,
                SEC_EST_MISSING,
                SEC_EST_SERVICE_CALL,
                SEC_EST_TBA,
                SEC_EST_SNAPSHOT,
                SEC_PENDING_REVIEW,
                *ests,
                SEC_INITIAL_UPLOADS,
                SEC_DAILY_UPLOADS,
                SEC_AUDIT_REJECTION,
                SEC_AUDIT_DISPUTE,
            ]
    except Exception:
        pass
    return list(_DEFAULT_SECTION_ORDER)


def _estimators_from_order(order):
    """Pull the estimator entries (non-builtin) out of a full section
    order in their original positions."""
    return [s for s in order if s not in _BUILTIN_SET]


# ── Franchise key ────────────────────────────────────────────────────────────
_FRANCHISE_PAREN_RE = re.compile(r"\s*\([^)]*\)")
# Trailing claim label — "1st Claim", "2nd Claim", "Claim 2", etc.
# Stripped for franchise lookup so multi-claim rows resolve the
# insured's one franchise tag. (Parentheticals are removed first.)
_FRANCHISE_CLAIM_RE = re.compile(
    r"\s*(?:\d+(?:st|nd|rd|th)?\s+claim|claim\s+\d+)\s*$", re.IGNORECASE)


def _franchise_key(text):
    s = (text or "").strip()
    if not s:
        return ""
    # Strip parentheticals like "(Contents)" / "(Self Pay)" — the
    # franchise belongs to the job, not the parenthetical variant.
    s = _FRANCHISE_PAREN_RE.sub("", s)
    # Normalize hyphen-spacing: "Smith- AAA", "Smith-AAA", "Smith - AAA"
    # all collapse to a single " - " separator before the split.
    s = re.sub(r"\s*-\s*", " - ", s)
    # Drop everything from the first " - " onward: that's the carrier /
    # sub / status tail, never part of the job identity for franchise
    # purposes. Bare names (no " - ") pass through unchanged.
    head = s.split(" - ", 1)[0]
    # Drop a trailing claim LABEL ("… 1st Claim", "… 2nd Claim (Kitchen)").
    # Franchise belongs to the INSURED, not the claim — both claims of a
    # multi-claim job share one franchise tag, so a per-claim APA row
    # ("Sayra Mansolino 2nd Claim") must still resolve the base tag.
    head = _FRANCHISE_CLAIM_RE.sub("", head)
    return " ".join(head.lower().split())


# ── Doc path for today ───────────────────────────────────────────────────────
WEEKDAY_SPELLING = {
    0: "Monday", 1: "Tuesday", 2: "Wednsday",  # file uses this spelling
    3: "Thursday", 4: "Friday", 5: "Saturday", 6: "Sunday",
}


def doc_path_for_today(dt=None):
    dt = dt or datetime.today()
    year  = dt.strftime("%Y")
    month = dt.strftime("%B")
    wd    = WEEKDAY_SPELLING[dt.weekday()]
    folder = os.path.join(_apa_root(), year, month)
    # File naming switched 2026-05-04 from "5-4-Monday .docx" (trailing
    # space before .docx) to plain "5-4-Monday.docx". The legacy April
    # archive uses the spaced form, so when the no-space file is missing
    # we still match a pre-existing spaced file rather than blindly
    # creating a duplicate. New writes always use the no-space form.
    no_space = os.path.join(folder, f"{dt.month}-{dt.day}-{wd}.docx")
    spaced   = os.path.join(folder, f"{dt.month}-{dt.day}-{wd} .docx")
    if os.path.isfile(no_space):
        return no_space
    if os.path.isfile(spaced):
        return spaced
    return no_space


# ── Item-text status/sub parsing ─────────────────────────────────────────────
# Sub-categories from Trello that don't have their own header in the APA doc.
SUB_OPTIONS = ["",
                "Initial Inspections/Re-Inspections",
                "Monitor",
                "Work in progress",
                "Upcoming/Pending",
                "TBS New Loss/Re-Inspection",
                "TBS Mitigation",
                "TBS Contents",
                "Testing/Clearance",
                "Office Questions",
                "PENDING APROVAL/INS/SELFPAY",
                "Pending Approvals/Property Management",
                "On Hold",
                "Initial",
                "Add'l Work/Missing Items"]

_FINAL_UPLOAD_EXTRAS = ["Completed - Work Performed"]
_INITIAL_DAILY_EXTRAS = ["TBA"]

# Trailing-string list shared by strip_status_from_text and gather_known_clients
# so adding a new SUB_OPTION automatically teaches both parsers to strip it.
# Status fragments stay listed explicitly — they don't live in SUB_OPTIONS.
_STATUS_TRAILINGS = (
    "-pending upload", "-pending(User)", "-pending(user)",
    "-Final upload", "-extended", "-uploaded", "-uploading",
    "-pending",
    "-not sold/cancelled", "-not sold", "-cancelled",
)
_SUB_TRAILINGS = tuple(
    f"-{s}" for s in SUB_OPTIONS + _FINAL_UPLOAD_EXTRAS + _INITIAL_DAILY_EXTRAS
    if s)
_ALL_TRAILINGS = _STATUS_TRAILINGS + _SUB_TRAILINGS

# Any status in this set highlights the row (yellow) in both GUI and saved doc
HIGHLIGHT_STATUSES = {"pending", "pending upload"}

# All status values that could appear in the doc (used for parsing/red coloring)
_ALL_STATUS_VARIANTS = ["pending upload", "pending(User)", "pending(user)",
                        "Final upload", "extended", "uploaded", "uploading", "pending",
                        "not sold/cancelled", "not sold", "cancelled"]


def strip_status_from_text(text):
    """Strip trailing -status / -sub from an item, leaving 'Client - Carrier'."""
    t = text.strip()
    for trailing in _ALL_TRAILINGS:
        if t.lower().endswith(trailing.lower()):
            t = t[:-len(trailing)].rstrip(" -—")
    return t


# ── Section-order cache + role groupings ─────────────────────────────────────
# Module-level cache so accessors don't hit persistence on every dropdown
# lookup. `_reload_estimators_cache` is called (by the web Manage-Sections
# action) after a saved reorder. apa_logic OWNS this cache for the web
# panels — apa_web / audit_web read it via module-attribute access
# (`apa.SECTION_ORDER`), so a reload is picked up immediately. The Tk
# apa_monitor_gui keeps its own copy of these mutable names because it
# reads them as locals and reloads them locally; a single running process
# is ever only web OR Tk, never both, so the two caches never cross.
SECTION_ORDER      = _persisted_section_order()
ESTIMATORS_ORDERED = _estimators_from_order(SECTION_ORDER)
ESTIMATOR_SECTIONS = set(ESTIMATORS_ORDERED)
SUB_SECTIONS       = {SEC_FINAL_UPLOADS, SEC_INITIAL_UPLOADS, SEC_DAILY_UPLOADS}
AUDIT_SECTIONS     = {SEC_AUDIT_REJECTION, SEC_AUDIT_DISPUTE}


def _reload_estimators_cache():
    """Refresh SECTION_ORDER + ESTIMATORS_ORDERED + ESTIMATOR_SECTIONS from
    persistence. Called after a Manage-Sections save. Name kept for
    backwards compatibility — it reloads the FULL section order now, not
    just the estimator slice."""
    global SECTION_ORDER, ESTIMATORS_ORDERED, ESTIMATOR_SECTIONS
    SECTION_ORDER = _persisted_section_order()
    ESTIMATORS_ORDERED = _estimators_from_order(SECTION_ORDER)
    ESTIMATOR_SECTIONS = set(ESTIMATORS_ORDERED)


# ── Dropdown option lists ────────────────────────────────────────────────────
# Default status dropdown (Final/Initial/Daily Uploads + estimator sections).
# "uploading" was replaced by "pending upload" (2026-06-17) so the
# highlight behavior is universal: "pending upload" is in
# HIGHLIGHT_STATUSES, so picking it turns the row yellow in every section
# exactly like the Audit Rejection/Dispute rows. Legacy "uploading" still
# parses (it's in _ALL_STATUS_VARIANTS) and wrap_item normalizes it to
# "pending upload".
STATUS_OPTIONS = ["", "pending", "extended", "uploaded", "pending upload",
                  "Final upload", "not sold/cancelled"]
# Status dropdown contents for Audit Rejection / Audit Dispute.
AUDIT_STATUS_OPTIONS = ["", "pending", "uploaded", "extended", "pending upload",
                        "not sold/cancelled"]


def _sub_options_for_section(section):
    """Return the list for the first (sub) dropdown, or None to hide it.

    Final/Initial/Daily Uploads share most of SUB_OPTIONS but each has a
    section-specific tail value: Final gets 'Completed - Work Performed'
    (only meaningful for closed-out jobs), Initial+Daily get 'TBA' (used
    when a job hasn't been assigned yet)."""
    # Audit Rejection/Dispute AND Pending Review use the ESTIMATOR as the
    # sub — a pending-review item belongs to whichever estimator owns it,
    # exactly like the dispute rows. (We don't fold PENDING REVIEW into
    # AUDIT_SECTIONS, which drives Teams rollups / note-appending; only
    # the sub dropdown should match.)
    if section in AUDIT_SECTIONS or section == SEC_PENDING_REVIEW:
        return [""] + ESTIMATORS_ORDERED
    if section == SEC_FINAL_UPLOADS:
        return SUB_OPTIONS + _FINAL_UPLOAD_EXTRAS
    if section in (SEC_INITIAL_UPLOADS, SEC_DAILY_UPLOADS):
        return SUB_OPTIONS + _INITIAL_DAILY_EXTRAS
    if section in SUB_SECTIONS:
        return SUB_OPTIONS
    return None


def _status_options_for_section(section):
    if section in AUDIT_SECTIONS:
        return AUDIT_STATUS_OPTIONS
    return STATUS_OPTIONS


# ── Item suffix parsers (shared by Tk + web) ─────────────────────────────────
def strip_to_base(text):
    """Reduce 'Last, First - Carrier-Initial-pending(User)' down to
    'Last, First - Carrier' so dedupe + note keys match across sub/status
    variants. Iterates because an item may carry BOTH a -sub and a -status
    trailing. Mirrors apa_monitor_gui._strip_to_base so web note keys line
    up with the ones the Tk panel writes."""
    base = re.sub(r'\s*-\s*\d+d\s+inactive.*$', '', (text or "").strip())
    while True:
        before = base
        for trailing in _ALL_TRAILINGS:
            if base.lower().endswith(trailing.lower()):
                base = base[:-len(trailing)].rstrip(" -—")
                break
        if base == before:
            return base


def wrap_item(text, highlighted=False):
    """Split 'Brew, Brian - AAA-Testing/Clearance-extended' into
    {text, sub, status, highlighted}. Estimator names count as subs (that's
    how Audit Rejection/Dispute rows record which estimator owns them), and
    plain `strip_status_from_text` does NOT peel estimator subs — so callers
    that need the sub MUST go through here. Longest known option wins so
    compound values aren't mis-parsed. Mirrors apa_monitor_gui._wrap_item
    (minus the UI-only franchise tag)."""
    item = {"text": text, "sub": "", "status": "",
            "highlighted": bool(highlighted)}
    remaining = (text or "").strip()

    # Peel status from the end (longest variant first).
    for s in sorted(_ALL_STATUS_VARIANTS, key=len, reverse=True):
        if remaining.lower().rstrip(".").endswith(s.lower()):
            idx = remaining.lower().rfind(s.lower())
            remaining = remaining[:idx].rstrip(" -—")
            if s.lower() == "pending upload":
                item["status"] = "pending upload"
            elif s.lower() == "uploading":
                # Legacy tag — normalize to the universal "pending upload"
                # so old entries pick up the highlight + the new label.
                item["status"] = "pending upload"
            elif s.lower().startswith("pending"):
                item["status"] = "pending"
            else:
                item["status"] = s
            break

    # Peel sub-category from the end (SUB_OPTIONS + estimator names).
    sub_candidates = [x for x in SUB_OPTIONS if x] + ESTIMATORS_ORDERED
    for s in sorted(sub_candidates, key=len, reverse=True):
        if remaining.lower().rstrip(".").endswith(s.lower()):
            idx = remaining.lower().rfind(s.lower())
            remaining = remaining[:idx].rstrip(" -—")
            item["sub"] = s
            break

    item["text"] = remaining
    return item


# ── Teams chat + estimator label ─────────────────────────────────────────────
def open_teams_chat(email, message):
    """Open the Teams desktop app to a chat with `email`, message pre-filled.
    User just hits Send. No API/auth required."""
    if not email:
        return False
    url = (f"msteams:/l/chat/0/0?users={urllib.parse.quote(email)}"
           f"&message={urllib.parse.quote(message)}")
    try:
        os.startfile(url)
        return True
    except OSError:
        return False


def estimator_first_name(section_name):
    """Convert 'JUAN' → 'Juan', 'AARON L' → 'Aaron L'."""
    return section_name.title()


# ── Doc parse / write ────────────────────────────────────────────────────────

def parse_existing_doc(path, section_order=None):
    """Parse an APA doc into {section_name: [(text, highlighted), ...]}.

    `section_order` controls which headers are recognized + the dict keys;
    defaults to the persisted order (== the old module-global SECTION_ORDER)."""
    order = section_order if section_order is not None else _persisted_section_order()
    sections = {s: [] for s in order}
    if not os.path.isfile(path):
        return sections
    try:
        doc = Document(path)
    except Exception as ex:
        try:
            import ems_log
            ems_log.warn("apa_monitor",
                f"could not open APA doc {path!r}: {ex}")
        except Exception:
            pass
        return sections

    def _match(text):
        t = text.strip().lower().rstrip(".").rstrip()
        for s in order:
            if t == s.lower().rstrip(".").rstrip():
                return s
        return None

    current = None
    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        m = _match(text)
        if m:
            current = m
            continue
        if not current:
            continue
        highlighted = any(
            getattr(run.font, "highlight_color", None) == WD_COLOR_INDEX.YELLOW
            for run in p.runs
        )
        # Strip leading "N. " numbering prefix so round-tripping stays clean
        text = _NUM_PREFIX_RE.sub("", text)
        sections[current].append((text, highlighted))
    return sections


def write_doc(path, today_date, sections, section_order=None):
    """Write the APA doc out with items numbered per section, smaller font,
    and '-extended' suffix rendered in red.

    `section_order` is the section sequence to emit; defaults to the
    persisted order (== the old module-global SECTION_ORDER)."""
    order = section_order if section_order is not None else _persisted_section_order()
    os.makedirs(os.path.dirname(path), exist_ok=True)
    doc = Document()
    doc.add_paragraph(today_date.strftime("%m/%d/%y"))

    def _add_run(paragraph, text, *, highlighted=False, red=False):
        run = paragraph.add_run(text)
        run.font.size = _ITEM_FONT_SIZE
        if red:
            run.font.color.rgb = _EXTENDED_RED
        if highlighted:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    for name in order:
        header_text = SEC_PENDING_REVIEW_DOC if name == SEC_PENDING_REVIEW else name
        doc.add_paragraph(header_text)

        # Build a filtered, ordered list so numbering skips blank items
        def _to_tuple(it):
            if isinstance(it, tuple):
                return it
            if isinstance(it, dict):
                return (
                    it.get("text", ""),
                    it.get("status", "").lower() in HIGHLIGHT_STATUSES,
                )
            return (str(it), False)

        items = [
            (t, h) for t, h in (_to_tuple(it) for it in sections.get(name, []))
            if t and t.strip()
        ]

        for idx, (text, highlighted) in enumerate(items, 1):
            p = doc.add_paragraph(style="List Paragraph")
            # Standard numbered-list indent: number at 0.25", body wraps at 0.5"
            p.paragraph_format.left_indent        = Inches(0.5)
            p.paragraph_format.first_line_indent  = Inches(-0.25)
            # Number prefix
            _add_run(p, f"{idx}. ", highlighted=highlighted)

            # Split out any trailing status (longest match wins) and color it red
            low = text.lower()
            status_found = None
            for s in sorted(_ALL_STATUS_VARIANTS, key=len, reverse=True):
                suffix = f"-{s.lower()}"
                if low.endswith(suffix):
                    status_found = s
                    break
            if status_found:
                suffix_len = len(status_found) + 1  # +1 for the '-'
                _add_run(p, text[:-suffix_len], highlighted=highlighted)
                _add_run(p, text[-suffix_len:], highlighted=highlighted, red=True)
            else:
                _add_run(p, text, highlighted=highlighted)

    doc.save(path)
