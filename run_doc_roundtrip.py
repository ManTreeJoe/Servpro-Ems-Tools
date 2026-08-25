"""Phase 0 round-trip harness for the live run doc.

See RUN_DOC_LIVE_PLAN.md §4. The assertion that matters:

    parse real .docx -> rows -> write .docx -> parse again
      => identical (jobs, run_date)

If that is not exact on real documents, the rest of the plan needs
rethinking, because the `.docx` stops being a faithful render of the rows
and printing silently drifts from what people typed.

Nothing here writes to `runs_dir`. Rendered copies go to a caller-supplied
output directory (the scratchpad by default) and real docs are only ever
read, with their mtimes restored (OneDrive Files-On-Demand stamps today's
mtime the first time it materializes a cloud-only file).

`write_run_doc` is a PROTOTYPE living here on purpose: phase 0 must not
change production `run_doc.py` while the two-machine trial is settling.
It moves into `run_doc.py` at phase 2.

Run:
    python run_doc_roundtrip.py                 # newest 10 docs in runs_dir
    python run_doc_roundtrip.py --all           # every doc in runs_dir
    python run_doc_roundtrip.py <path.docx> ... # specific docs
"""
from __future__ import annotations

import os
import re
import sys
import glob
import tempfile

import config
import run_doc


# --- rows ------------------------------------------------------------------
#
# A row is one Word paragraph, stored AS AUTHORED (plan §1). `kind` is a
# derived read-time view used only by the renumbering variant below — it is
# never the source of truth for the text.

_STOP_RE = re.compile(r'^(upcoming|tbs\b|pending|on hold|marketing)', re.IGNORECASE)
_NUM_RE = re.compile(r'^(\d+)\.\s*')


def doc_to_rows(path):
    """Read a run doc into rows: [{'text', 'struck', 'kind', 'section'}].

    Every paragraph becomes a row, including blanks, headings and lines the
    parser ignores, so the writer can reproduce the document in full. Blank
    paragraphs are kept because they carry the page's visual spacing.
    """
    snap = run_doc._preserve_mtime(path)
    from docx import Document
    from audit_logic import para_is_struck
    doc = Document(path)
    raw = [(p.text, para_is_struck(p)) for p in doc.paragraphs]
    run_doc._restore_mtime(path, snap)

    rows = []
    section = None
    for text, struck in raw:
        stripped = (text or "").strip()
        tl = stripped.lower()
        if not stripped:
            kind = "blank"
        elif "work to be performed" in tl:
            kind, section = "section", "work"
        elif re.match(r'^monitor', tl):
            kind, section = "section", "monitor"
        elif _STOP_RE.match(tl):
            kind, section = "stop", None
        elif section and not struck and not re.search(r'\bwarehouse\b', tl):
            kind = "entry"
        else:
            kind = "other"
        rows.append({"text": text or "", "struck": bool(struck),
                     "kind": kind, "section": section})
    return rows


# --- writer (prototype) ----------------------------------------------------

def _replace_paragraph_text(para, text):
    """Replace a paragraph's text while retaining its first run formatting.

    Most rows are unchanged during an export, so leave those paragraphs
    completely alone.  For an edited/renumbered row, keep the formatting of
    the first text-bearing run and clear the others.  This preserves the
    document template's font, size, bold, spacing and list geometry instead
    of falling back to Word's Normal defaults.
    """
    if para.text == text:
        return
    runs = list(para.runs)
    if not runs:
        para.add_run(text)
        return
    target = next((run for run in runs if run.text), runs[0])
    target.text = text
    for run in runs:
        if run is not target:
            run.text = ""


def write_run_doc(path, rows, *, renumber=False, template_path=None):
    """Render rows back to a `.docx`.

    Each row writes its text verbatim, with strikethrough when `struck` —
    strikethrough is MEANING here (the parser skips struck lines), not
    styling, so it has to survive the trip.

    `template_path` is the current printable document.  Opening it and
    changing only paragraph contents preserves the page setup, headers,
    styles and run formatting.  That is required while Word remains the
    live printable view; creating `Document()` from scratch changes a real
    three-page run into five pages even though its parsed jobs still match.

    `renumber=True` regenerates the "1. 2. 3." prefixes per section on the
    way out, as plan §4 step 2 describes. That is the variant under test:
    it is only safe if the numbers it produces match what was authored.
    """
    from docx import Document
    doc = Document(template_path) if template_path else Document()
    paras = list(doc.paragraphs)
    if template_path and len(paras) != len(rows):
        raise ValueError(
            f"template has {len(paras)} paragraphs but rows has {len(rows)}")
    counters = {}
    for index, row in enumerate(rows):
        text = row["text"]
        if renumber and row["kind"] == "entry" and _NUM_RE.match(text.strip()):
            sec = row.get("section") or ""
            counters[sec] = counters.get(sec, 0) + 1
            body = _NUM_RE.sub('', text.strip())
            text = f"{counters[sec]}. {body}"
        para = paras[index] if template_path else doc.add_paragraph()
        _replace_paragraph_text(para, text)
        runs = [run for run in para.runs if run.text]
        if not runs and text:
            runs = [para.add_run(text)]
        # Only SET strike, never clear it. `run.font.strike = False` makes
        # python-docx emit `<w:strike w:val="0"/>`, and audit_logic.
        # para_is_struck treats any strike element whose val is not the
        # literal string "false" as struck — so writing False would mark
        # every line struck and the parser would skip the whole document.
        # Leaving the attribute unset omits the element entirely.
        if row["struck"]:
            for run in runs:
                run.font.strike = True
    doc.save(path)
    return path


# --- comparison ------------------------------------------------------------

def _diff_jobs(a, b):
    """First few differences between two parsed job lists, as readable text."""
    out = []
    if len(a) != len(b):
        out.append(f"job count {len(a)} -> {len(b)}")
    for i, (ja, jb) in enumerate(zip(a, b)):
        if ja == jb:
            continue
        keys = sorted(set(ja) | set(jb))
        for k in keys:
            if ja.get(k) != jb.get(k):
                out.append(f"[{i}] {k}: {ja.get(k)!r} -> {jb.get(k)!r}")
        if len(out) > 12:
            break
    return out[:12]


def roundtrip(path, out_dir, *, renumber=False):
    """Parse -> rows -> write -> parse. Returns (ok, detail_lines)."""
    jobs_a, date_a = run_doc.parse_run_doc(path)
    rows = doc_to_rows(path)
    tmp = os.path.join(out_dir, ("renum_" if renumber else "verbatim_")
                       + os.path.basename(path).replace(" ", "_"))
    write_run_doc(tmp, rows, renumber=renumber, template_path=path)
    jobs_b, date_b = run_doc.parse_run_doc(tmp)

    detail = []
    if date_a != date_b:
        detail.append(f"run_date {date_a!r} -> {date_b!r}")
    detail += _diff_jobs(jobs_a, jobs_b)
    return (not detail), detail, len(jobs_a), len(rows)


# --- cli -------------------------------------------------------------------

def _discover(limit=None):
    runs = config.load().get("runs_dir") or ""
    if not runs or not os.path.isdir(runs):
        return []
    found = [p for p in glob.glob(os.path.join(runs, "*", "*.docx"))
             if not os.path.basename(p).startswith("~$")]
    found.sort(key=lambda p: os.path.getmtime(p), reverse=True)
    return found[:limit] if limit else found


def main(argv):
    args = [a for a in argv[1:] if not a.startswith("--")]
    take_all = "--all" in argv
    paths = args or _discover(None if take_all else 10)
    if not paths:
        print("no run docs found")
        return 1

    out_dir = os.path.join(tempfile.gettempdir(), "run_doc_roundtrip")
    os.makedirs(out_dir, exist_ok=True)

    fails = {"verbatim": 0, "renumber": 0}
    print(f"{len(paths)} doc(s); rendered copies in {out_dir}\n")
    for p in paths:
        name = os.path.basename(p)
        line = [f"{name[:34]:34}"]
        for mode in ("verbatim", "renumber"):
            try:
                ok, detail, njobs, nrows = roundtrip(
                    p, out_dir, renumber=(mode == "renumber"))
            except Exception as exc:
                fails[mode] += 1
                line.append(f"{mode}=ERR({exc.__class__.__name__})")
                continue
            if not ok:
                fails[mode] += 1
            line.append(f"{mode}={'PASS' if ok else 'FAIL'}")
            if not ok:
                line.append("\n" + "\n".join("      - " + d for d in detail))
        print("  ".join(line))

    print(f"\nverbatim: {len(paths)-fails['verbatim']}/{len(paths)} pass"
          f"   renumber: {len(paths)-fails['renumber']}/{len(paths)} pass")
    return 0 if fails["verbatim"] == 0 else 2


if __name__ == "__main__":
    raise SystemExit(main(sys.argv))
