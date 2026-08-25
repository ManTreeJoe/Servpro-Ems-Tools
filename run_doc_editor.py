"""Safe editing core for every operational section of a daily run document.

The Word file remains the source of truth. Recognised dispatch sections are
editable here; every unrecognised paragraph remains untouched.

Every save is conflict-checked, backed up, written to a sibling temporary
file, validated by reopening it, and atomically replaced.  This module has
no webview dependency so the document rules can be tested in isolation.
"""
from __future__ import annotations

import copy
import datetime as _dt
import hashlib
import os
import re
import shutil
import uuid

import paths


SECTION_DEFS = (
    ("monitor", "Monitor", r"^monitor\b"),
    ("work", "Work to be performed", r"^work to be performed\b"),
    ("upcoming", "Upcoming", r"^upcoming\b"),
    ("tbs_new_loss", "TBS New Loss / Reinspection", r"^tbs new loss\b"),
    ("tbs_mitigation", "TBS Mitigation", r"^tbs mitigation\b"),
    ("tbs_contents", "TBS Contents", r"^tbs contents\b"),
    ("pending_testing", "Pending Testing / Clearance / Abatement",
     r"^pending testing"),
    ("pending_insurance", "Pending Approvals — Insurance / Self Pay",
     r"^pending approvals.*insurance"),
    ("pending_property", "Pending Approvals — Property Management",
     r"^pending approvals.*property"),
    ("on_hold", "On Hold", r"^on hold\b"),
    ("marketing", "Marketing Team — Insurance / Self Pay", r"^marketing team"),
)
SECTIONS = tuple(item[0] for item in SECTION_DEFS)
SECTION_LABELS = {key: label for key, label, _pattern in SECTION_DEFS}
_SECTION_PATTERNS = tuple((key, re.compile(pattern, re.IGNORECASE))
                          for key, _label, pattern in SECTION_DEFS)
_MAX_ROWS = 200
_MAX_TEXT = 2000


class RunDocConflict(RuntimeError):
    """The file changed after the editor loaded it."""


def file_version(path: str) -> str:
    digest = hashlib.sha256()
    with open(path, "rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _section_for(text: str, current: str | None) -> str | None:
    probe = (text or "").strip().lower()
    for key, pattern in _SECTION_PATTERNS:
        if pattern.search(probe):
            return key
    return current


def _paragraph_is_struck(paragraph) -> bool:
    from audit_logic import para_is_struck
    return bool(para_is_struck(paragraph))


def _scan(document):
    """Return section headings and editable paragraph slots."""
    headings = {}
    slots = {name: [] for name in SECTIONS}
    current = None
    for paragraph in document.paragraphs:
        text = paragraph.text or ""
        previous = current
        current = _section_for(text, current)
        if current in SECTIONS and current != previous:
            headings[current] = paragraph
            continue
        if current in SECTIONS and text.strip():
            # Warehouse and struck rows are deliberately included. They are
            # dispatch lines even though the audit parser ignores them.
            slots[current].append(paragraph)
    return headings, slots


def read_document(path: str) -> dict:
    from docx import Document

    document = Document(path)
    headings, slots = _scan(document)
    present = [name for name in SECTIONS if name in headings]
    sections = {}
    for name in present:
        sections[name] = [
            {
                "id": f"{name}:{index}",
                "text": paragraph.text or "",
                "struck": _paragraph_is_struck(paragraph),
            }
            for index, paragraph in enumerate(slots[name])
        ]
    stat = os.stat(path)
    return {
        "path": os.path.abspath(path),
        "filename": os.path.basename(path),
        "version": file_version(path),
        "modified": _dt.datetime.fromtimestamp(stat.st_mtime).isoformat(
            timespec="seconds"),
        "sections": sections,
        "section_order": present,
        "section_labels": {name: SECTION_LABELS[name] for name in present},
    }


def _normalise_rows(rows) -> list[dict]:
    if not isinstance(rows, list):
        raise ValueError("section rows must be a list")
    if len(rows) > _MAX_ROWS:
        raise ValueError(f"a section cannot contain more than {_MAX_ROWS} rows")
    clean = []
    for row in rows:
        if not isinstance(row, dict):
            raise ValueError("each row must be an object")
        text = str(row.get("text") or "")
        if "\r" in text or "\n" in text:
            raise ValueError("each run-doc row must stay on one line")
        if not text.strip():
            continue
        if len(text) > _MAX_TEXT:
            raise ValueError(f"a row cannot exceed {_MAX_TEXT} characters")
        clean.append({"text": text, "struck": bool(row.get("struck"))})
    return clean


def _replace_text(paragraph, text: str) -> None:
    if paragraph.text == text:
        return
    runs = list(paragraph.runs)
    if not runs:
        paragraph.add_run(text)
        return
    target = next((run for run in runs if run.text), runs[0])
    target.text = text
    for run in runs:
        if run is not target:
            run.text = ""


def _set_strike(paragraph, struck: bool) -> None:
    """Set or remove strike XML without writing the ambiguous val=0 form."""
    from docx.oxml.ns import qn

    runs = [run for run in paragraph.runs if run.text]
    if not runs and paragraph.text:
        runs = [paragraph.add_run(paragraph.text)]
    for run in runs:
        props = run._r.get_or_add_rPr()
        for node in list(props):
            if node.tag == qn("w:strike"):
                props.remove(node)
        if struck:
            run.font.strike = True


def _new_slot(document, heading, template, after):
    """Create an entry paragraph after ``after``, cloning local styling."""
    from docx.text.paragraph import Paragraph

    if template is not None:
        element = copy.deepcopy(template._p)
        # Content is replaced immediately, but clearing it here prevents a
        # duplicate line if an exception interrupts the edit.
        for node in element.xpath(".//w:t"):
            node.text = ""
    else:
        scratch = document.add_paragraph()
        element = scratch._p
        element.getparent().remove(element)
    anchor = after._p if after is not None else heading._p
    anchor.addnext(element)
    return Paragraph(element, document._body)


def _write_section(document, name: str, desired: list[dict]) -> None:
    headings, all_slots = _scan(document)
    heading = headings.get(name)
    if heading is None:
        raise ValueError(f"document has no {name.title()} section")
    slots = all_slots[name]
    template = slots[-1] if slots else None

    while len(slots) < len(desired):
        after = slots[-1] if slots else heading
        slots.append(_new_slot(document, heading, template, after))

    for paragraph, row in zip(slots, desired):
        _replace_text(paragraph, row["text"])
        _set_strike(paragraph, row["struck"])

    # Delete surplus entry slots from the end. Static paragraphs below the
    # section are separate XML nodes and stay exactly where they were.
    for paragraph in reversed(slots[len(desired):]):
        element = paragraph._p
        element.getparent().remove(element)


def _backup(path: str) -> str:
    dest = paths.data("run_doc_backups")
    os.makedirs(dest, exist_ok=True)
    stamp = _dt.datetime.now().strftime("%Y%m%d-%H%M%S")
    base = os.path.basename(path)
    target = os.path.join(dest, f"{stamp}__{base}")
    shutil.copy2(path, target)
    copies = sorted(
        (os.path.join(dest, name) for name in os.listdir(dest)
         if name.endswith("__" + base)),
        key=os.path.getmtime,
        reverse=True,
    )
    for old in copies[20:]:
        try:
            os.remove(old)
        except OSError:
            pass
    return target


def save_document(path: str, expected_version: str, sections: dict) -> dict:
    """Save Monitor/Work rows while preserving every other paragraph."""
    if not os.path.isfile(path) or not path.lower().endswith(".docx"):
        raise ValueError("an existing Word .docx is required")
    if file_version(path) != (expected_version or ""):
        raise RunDocConflict(
            "The Word file changed after this screen loaded. Reload it and "
            "review the newer copy before saving.")
    if not isinstance(sections, dict):
        raise ValueError("sections are required")
    requested = [name for name in SECTIONS if name in sections]
    desired = {name: _normalise_rows(sections.get(name, []))
               for name in requested}

    from docx import Document
    document = Document(path)
    for name in requested:
        _write_section(document, name, desired[name])

    backup = _backup(path)
    folder = os.path.dirname(path)
    temp = os.path.join(
        folder, f".{os.path.basename(path)}.linguar-{uuid.uuid4().hex[:8]}.docx")
    try:
        document.save(temp)
        check = read_document(temp)
        for name in requested:
            got = [{"text": row["text"], "struck": row["struck"]}
                   for row in check["sections"][name]]
            if got != desired[name]:
                raise ValueError(f"saved {name} rows did not validate")
        os.replace(temp, path)
    except Exception:
        try:
            if os.path.isfile(temp):
                os.remove(temp)
        except OSError:
            pass
        raise

    result = read_document(path)
    result["backup"] = backup
    return result
